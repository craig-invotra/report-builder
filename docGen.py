from docx import Document
import dataHandle as dh
from dataHandle import maketable
from dataHandle import makefiltertable
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import os
import subprocess

#create document
document = Document()

#set style
style = document.styles['Normal']
font = style.font
font.name = 'Ubuntu'
font.size = Pt(10)

title = document.styles['Title']
T_font = title.font
T_font.name = 'Ubuntu'
T_font.size = Pt(26)

head2 = document.styles['Heading 2']
hfont = head2.font
hfont.name = 'Ubuntu'
hfont.size = Pt(14)

new_heading_style = document.styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = document.styles['Heading 2']
font2 = new_heading_style.font
font2.name = 'Ubuntu'
font2.size = Pt(16)
#font2.color.ColorFormat

#user inputs for file genearation
productVers = raw_input("Please enter the Product version: ")
department = raw_input("Please enter the Client name: ")
CVSfile = "JIRA.csv"

tableData= dh.infostrip(CVSfile)
docTitle = productVers + ' Testing Results Report ' + department + '- UAT'

#generate the body of the document
document.add_heading(docTitle, 0)

document.add_heading('1 Test Summary', level=2,)
document.add_paragraph('The following tests have been run on ' + department + 's Live environment.')
t1 = document.add_table(rows=4, cols=3)
tlc = t1.rows[0].cells
tlc[0].text = 'Total JIRA tasks'
tlc[1].text = str(len(tableData))
tlc[2].text = 'No. of tasks in the release as a whole'
tlc2 = t1.rows[1].cells
tlc2[0].text = 'Verified UAT'
tlc2[1].text = str( len( tableData[tableData['Status'] == 'Verified - UAT'] ) )
tlc2[2].text = 'Issues that have been assigned over to DFT to carry out testing'
tlc3 = t1.rows[2].cells
tlc3[0].text = 'Reopened'
tlc3[1].text = str( len( tableData[tableData['Status'] == 'Reopened'] ) )
tlc3[2].text = 'Issues have gone back to the development team for investigation Refer to section 4.2'
tlc4 = t1.rows[3].cells
tlc4[0].text = 'Closed'
tlc4[1].text = str( len( tableData[tableData['Status'] == 'Closed'] ) )
tlc4[2].text = 'Issues that do not require the client to test. E.g. functionality that is disabled or ' \
               'maintenance improvements'


document.add_heading('1.1 Breakdown of JIRA tests', level=2)
document.add_paragraph('')
t2 = document.add_table(rows=1, cols=4)
t2c = t2.rows[0].cells
t2c[0].text = "Name"
t2c[1].text = "Summary"
t2c[2].text = "Status"
t2c[3].text = "Pass"
maketable(tableData, t2)



document.add_page_break()

document.add_heading('2 Regression Tests', level=2)
document.add_paragraph('The table below summarises the tests run for regression testing and the test results obtained for each regression test:')

makefiltertable(document, tableData, 'Closed')

document.add_page_break()

document.add_heading('3 Performance Testing', level=2)
document.add_paragraph('The table below summarises the tests run for performance testing and the test results obtained for each test:')

document.add_page_break()

document.add_heading('4 Test Instances', level=2)
document.add_paragraph('This sections provides information on any unexpected results, problems or defects that occurred during testing. All defects will be logged in JIRA for reference and retesting when required. ')

document.add_heading('4.1 Resolved Test Instances', level=2)
document.add_paragraph('It is known that bugs can arise from testing post deployment, this may be due to deployment issue or bugs missed during QA due to environmental differences. All defects will be logged in JIRA and assessed and progressed appropriately. ')
makefiltertable(document, tableData, 'Closed')

document.add_heading('4.2 Unresolved Test Instances', level=2)
document.add_paragraph('Unfortunately it may be the case that some defects cannot be resolved within the current release for various reasons. If this is the case,  a plan of next steps will be taken and adhered to. ')


#save the document object

dire =os.getenv("HOME") +'/Documents/Generated Reports/'

if not os.path.exists(dire):
    os.makedirs(dire)
document.save(dire + docTitle + '.docx')
subprocess.Popen(["xdg-open", dire])

